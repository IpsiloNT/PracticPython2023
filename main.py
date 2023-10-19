import json
from prettytable import PrettyTable
from datetime import datetime
from docx import Document
from docx2pdf import convert
import pandas as pd
import os
import msvcrt
import numpy as np
import matplotlib.pyplot as plt
import atexit
import calendar
import datetime
import calendar
import datetime
import numpy as np


JSON_FILE = "users.json"

def load_data():
    try:
        with open(JSON_FILE, "r") as f:
            data = json.load(f)
    except (FileNotFoundError, json.decoder.JSONDecodeError):
        data = []
    return data

# Функция для сохранения данных в JSON-файл
def save_data(data):
    with open(JSON_FILE, "w") as f:
        json.dump(data, f, indent=4)

# Функция для проверки существования пользователя по логину
def check_user_existence(login, data):
    return any(user["login"] == login for user in data)

# Функция для авторизации пользователя
def authenticate(data):
    while True:
        user_login = input("\033[1;32mВведите логин: \033[0m")

        password = ""
        password = password.encode('iso-8859-1').decode('utf-8')
        print('\033[1;32mВведите пароль: ', end='')
        while True:
            ch = msvcrt.getch()
            if ch == b'\r':
                print()
                break
            elif ch == b'\x08':
                if len(password) > 0:
                    password = password[:-1]
                    print('\b \b', end='')
            else:
                password += ch.decode("utf-8")
                print('\033[1;35m*', end='')

        user = next((user for user in data if user["login"] == user_login and user["password"] == password), None)

        if user:
            if user["status"] == "inactive":
                print("Вы не можете войти, так как вы отключены от системы. Обратитесь к администратору.")
                continue

            print(f"Вы вошли как {'Администратор' if user['role'] == 1 else 'Пользователь'}")
            user["login_count"] += 1
            user["login_time"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            user["logout_time"] = None
            save_data(data)

            if user["role"] == 1:
                admin_menu(data, user_login)
            elif user["role"] == 0:
                user_menu(data, user_login)

        else:
            print("Ошибка авторизации. Проверьте логин и пароль.")

def logout_user(data, user_login):
    user = next((user for user in data if user["login"] == user_login), None)
    if user:
        user["logout_time"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        save_data(data)


# Функция для отображения данных всех пользователей через PrettyTable
def view_users(data):
    if not data:
        print("Нет зарегистрированных пользователей.")
    else:
        table = PrettyTable()
        table.field_names = ["ID", "Фамилия", "Имя", "Логин", "Пароль", "Роль", "Статус"]

        for user in data:
            role = "Администратор" if user["role"] == 1 else "Пользователь"
            status = "Включен" if user["status"] == "active" else "Отключен"
            table.add_row([user["id"], user["surname"], user["name"], user["login"], user["password"], role, status])
        print(table)

# Функция для добавления нового пользователя
def add_user(data):
    last_name = input("Введите фамилию: ")
    first_name = input("Введите имя: ")

    while True:
        login = input("Введите логин: ")
        if not login:
            print("Логин не может быть пустым. Пожалуйста, введите логин.")
        elif check_user_existence(login, data):
            print("Пользователь с таким логином уже существует. Пожалуйста, выберите другой логин.")
        else:
            break

    while True:
        password = input("Введите пароль: ")
        if len(password) < 6:
            print("Пароль должен содержать как минимум 6 символов. Пожалуйста, введите пароль заново.")
        else:
            break

    while True:
        role = input("Введите роль (0 для пользователя, 1 для администратора): ")
        if role not in ("0", "1"):
            print("Роль должна быть 0 или 1.")
        else:
            role = int(role)
            break

    user_id = len(data) + 1
    user = {
        "id": user_id,
        "surname": last_name,
        "name": first_name,
        "login": login,
        "password": password,
        "role": role,
        "status": "active"
    }

    data.append(user)
    save_data(data)
    print("Пользователь успешно добавлен в базу данных.")
    view_users(data)

# Функция для изменения статуса пользователя
def change_user_status(data):
    while True:
        view_users(data)
        login_to_change_status = input("Введите логин пользователя, статус которого вы хотите изменить (или нажмите Enter для отмены): ")

        if not login_to_change_status:
            break

        user_to_change_status = next((user for user in data if user["login"] == login_to_change_status), None)

        if user_to_change_status:
            if user_to_change_status["status"] == "active":
                user_to_change_status["status"] = "inactive"
                print(f"Статус пользователя с логином '{login_to_change_status}' изменен на 'Отключен'.")
            else:
                user_to_change_status["status"] = "active"
                print(f"Статус пользователя с логином '{login_to_change_status}' изменен на 'Включен'.")
            save_data(data)
        else:
            print(f"Пользователь с логином '{login_to_change_status}' не найден.")

# Функция для удаления пользователя
def delete_user(data):
    view_users(data)
    login_to_delete = input("Введите логин пользователя, которого вы хотите удалить: ")

    user_to_delete = next((user for user in data if user["login"] == login_to_delete), None)

    if user_to_delete:
        data.remove(user_to_delete)
        save_data(data)
        print(f"Пользователь с логином '{login_to_delete}' успешно удален.")
        view_users(data)
    else:
        print(f"Пользователь с логином '{login_to_delete}' не найден.")

# Функция для изменения данных пользователя
def change_user_data(data):
    view_users(data)
    login_to_change = input("Введите логин пользователя, данные которого вы хотите изменить (или нажмите Enter для отмены): ")

    user_to_change = next((user for user in data if user["login"] == login_to_change), None)

    if user_to_change:
        print("Введите новые данные для пользователя (оставьте поле пустым, чтобы оставить без изменений):")
        new_last_name = input(f"Новая фамилия ({user_to_change['surname']}): ")
        new_first_name = input(f"Новое имя ({user_to_change['name']}): ")
        new_login = input(f"Новый логин ({user_to_change['login']}): ")

        while True:
            if check_user_existence(new_login, data):
                print("Пользователь с таким логином уже существует.")
                new_login = input("Пожалуйста, введите другой логин: ")
            else:
                break

        while True:
            new_password = input(f"Новый пароль ({user_to_change['password']}): ")
            if new_password and len(new_password) < 6:
                print("Пароль должен содержать как минимум 6 символов.")
            else:
                break

        if new_last_name:
            user_to_change['surname'] = new_last_name
        if new_first_name:
            user_to_change['name'] = new_first_name
        if new_login:
            user_to_change['login'] = new_login
        if new_password:
            user_to_change['password'] = new_password

        save_data(data)
        print(f"Данные пользователя с логином '{login_to_change}' успешно изменены.")
        view_users(data)
    elif login_to_change:
        print(f"Пользователь с логином '{login_to_change}' не найден.")

# Функция для показа опции сортировки
def print_sorted_surname(data):
    table = PrettyTable()
    table.field_names = ["Фамилия"]

    for user in data:
        table.add_row([user["surname"]])

    print(table)

def print_sorted_name(data):
    table = PrettyTable()
    table.field_names = ["Имя"]

    for user in data:
        table.add_row([user["name"]])
    print(table)

def print_sorted_login(data):
    table = PrettyTable()
    table.field_names = ["Логин"]

    for user in data:
        table.add_row([user["login"]])

    print(table)

# Функция для показа опции сортировки
def sort_menu(data):
    while True:
        print("\033[96mСортировка:\033[0m")
        print("1. По Фамилии")
        print("2. По Имени")
        print("3. По Логину")
        print("4. Назад")
        sort_choice = input("Выберите опцию сортировки: ")

        if sort_choice == "1":
            while True:
                print("\nСортировка по фамилии:")
                print("1. По возрастанию")
                print("2. По убыванию")
                print("3. Назад в меню сортировки")
                order_choice = input("Выберите порядок сортировки: ")

                if order_choice == "1":
                    data.sort(key=lambda user: user["surname"])
                    print_sorted_surname(data)
                elif order_choice == "2":
                    data.sort(key=lambda user: user["surname"], reverse=True)
                    print_sorted_surname(data)
                elif order_choice == "3":
                    break
                else:
                    print("Неверный выбор порядка сортировки. Пожалуйста, выберите 1, 2 или 3.")
        elif sort_choice == "2":
            while True:
                print("\nСортировка по имени:")
                print("1. По возрастанию")
                print("2. По убыванию")
                print("3. Назад в меню сортировки")
                order_choice = input("Выберите порядок сортировки: ")

                if order_choice == "1":
                    data.sort(key=lambda user: user["name"])
                    print_sorted_name(data)
                elif order_choice == "2":
                    data.sort(key=lambda user: user["name"], reverse=True)
                    print_sorted_name(data)
                elif order_choice == "3":
                    break
                else:
                    print("Неверный выбор порядка сортировки. Пожалуйста, выберите 1, 2 или 3.")
        elif sort_choice == "3":
            while True:
                print("\nСортировка по логину:")
                print("1. По возрастанию")
                print("2. По убыванию")
                print("3. Назад в меню сортировки")
                order_choice = input("Выберите порядок сортировки: ")

                if order_choice == "1":
                    data.sort(key=lambda user: user["login"])
                    print_sorted_login(data)
                elif order_choice == "2":
                    data.sort(key=lambda user: user["login"], reverse=True)
                    print_sorted_login(data)
                elif order_choice == "3":
                    break
                else:
                    print("Неверный выбор порядка сортировки. Пожалуйста, выберите 1, 2 или 3.")
        elif sort_choice == "4":
            break
        else:
            print("Неверный выбор сортировки. Пожалуйста, выберите 1, 2, 3 или 4.")

def filter_active_users(data):
    print("\033[96mАктивные пользователи:\033[0m")
    active_users = [user for user in data if user['status'] == 'active']
    if not active_users:
        print("Нет активных пользователей.")
    else:
        display_users1(active_users)

def filter_inactive_users(data):
    print("\033[96mНеактивные пользователи:\033[0m")
    inactive_users = [user for user in data if user['status'] == 'inactive']
    if not inactive_users:
        print("Нет неактивных пользователей.")
    else:
        display_users1(inactive_users)


def filter_admins(data):
    print("\033[96mАдминистраторы:\033[0m")
    admins = [user for user in data if user['role'] == 1]
    if not admins:
        print("Нет администраторов.")
    else:
        display_users2(admins)

def filter_users(data):
    print("\033[96mПользователи:\033[0m")
    users = [user for user in data if user['role'] == 0]
    if not users:
        print("Нет пользователей.")
    else:
        display_users2(users)

# Функция для поиска данных по введенному слову
def search_data_by_word(data, search_query):
    results = []

    for user in data:
        user_info = [str(user["id"]), user["surname"], user["name"], user["login"], user["password"],
                     "Администратор" if user["role"] == 1 else "Пользователь",
                     "Включен" if user["status"] == "active" else "Отключен"]

        for field in user_info:
            if search_query in field:
                results.append(user_info)
                break

    return results


def search_by_word(data):
    search_query = input("Введите слово для поиска: ")
    results = search_data_by_word(data, search_query)

    if results:
        table = PrettyTable()
        table.field_names = ["ID", "Фамилия", "Имя", "Логин", "Пароль", "Роль", "Статус"]

        for user_info in results:
            table.add_row(user_info)

        print("Найдены пользователи, чьи атрибуты содержат введенное слово:")
        print(table)
    else:
        print("Нет совпадений.")


def display_users1(users):
    # В этой функции можно отобразить информацию о пользователях (например, используя PrettyTable)
    # Примерно так:
    table = PrettyTable()
    table.field_names = ["ID", "Фамилия", "Имя", "Логин", "Статус"]
    for user in users:
        table.add_row([user['id'], user['surname'], user['name'], user['login'], user['status']])
    print(table)

def display_users2(users):
    # В этой функции можно отобразить информацию о пользователях (например, используя PrettyTable)
    # Примерно так:
    table = PrettyTable()
    table.field_names = ["ID", "Фамилия", "Имя", "Логин", "Роль"]
    for user in users:
        role = "Администратор" if user["role"] == 1 else "Пользователь"
        table.add_row([user['id'], user['surname'], user['name'], user['login'], role])
    print(table)


def filter_menu(data):
    while True:
        print("\033[96mФильтрация:\033[0m")
        print("1. Активные пользователи")  # Опция для фильтрации активных пользователей
        print("2. Отключенные пользователи")  # Опция для фильтрации отключенных пользователей
        print("3. Администраторы")  # Опция для фильтрации администраторов
        print("4. Пользователи")  # Опция для фильтрации обычных пользователей
        print("5. Назад")  # Опция для выхода из меню фильтрации
        filter_choice = input("Выберите опцию фильтрации: ")

        if filter_choice == "1":
            filter_active_users(data)
        elif filter_choice == "2":
            filter_inactive_users(data)
        elif filter_choice == "3":
            filter_admins(data)
        elif filter_choice == "4":
            filter_users(data)
        elif filter_choice == "5":
            break  # Выход из меню фильтрации
        else:
            print("Неверный выбор. Пожалуйста, выберите 1, 2, 3, 4 или 5.")


# Функция для выбора сортировка/фильтрация
def sort_or_filter(data):
    while True:
        print("1. Сортировка")
        print("2. Фильтрация")
        print("3. Поиск")
        print("4. Назад")
        option_choice = input("Выберите опцию: ")

        if option_choice == "1":
            sort_menu(data)
        elif option_choice == "2":
            filter_menu(data)
        elif option_choice == "3":
            search_by_word(data)
        elif option_choice == "4":
            break
        else:
            print("Неверный выбор. Пожалуйста, выберите 1, 2, 3 или 4.")


def print_sorted_surname(data):
    table = PrettyTable()
    table.field_names = ["Фамилия"]

    for user in data:
        table.add_row([
            user["surname"]
        ])

    print(table)

def print_sorted_name(data):
    table = PrettyTable()
    table.field_names = ["Имя"]

    for user in data:
        table.add_row([
            user["name"]
        ])

    print(table)

def print_sorted_login(data):
    table = PrettyTable()
    table.field_names = ["Логин"]

    for user in data:
        table.add_row([
            user["login"]
        ])

    print(table)

# Функция для показа опции сортировки
def sort_menu():
    while True:
        print("\033[96mСортировка:\033[0m")
        print("1. По Фамилии")
        print("2. По Имени")
        print("3. По Логину")
        print("4. Назад")
        sort_choice = input("Выберите опцию сортировки: ")

        if sort_choice == "1":
            while True:
                print("\nСортировка по фамилии:")
                print("1. По возрастанию")
                print("2. По убыванию")
                print("3. Назад в меню сортировки")
                order_choice = input("Выберите порядок сортировки: ")

                if order_choice == "1":
                    data.sort(key=lambda user: user["surname"])
                    print_sorted_surname(data)
                    pass
                elif order_choice == "2":
                    data.sort(key=lambda user: user["surname"], reverse=True)
                    print_sorted_surname(data)
                    pass
                elif order_choice == "3":
                    break
                else:
                    print("Неверный выбор порядка сортировки. Пожалуйста, выберите 1, 2 или 3.")
        elif sort_choice == "2":
            while True:
                print("\nСортировка по имени:")
                print("1. По возрастанию")
                print("2. По убыванию")
                print("3. Назад в меню сортировки")
                order_choice = input("Выберите порядок сортировки: ")

                if order_choice == "1":
                    data.sort(key=lambda user: user["name"])
                    print_sorted_name(data)
                    pass
                elif order_choice == "2":
                    data.sort(key=lambda user: user["name"], reverse=True)
                    print_sorted_name(data)
                    pass
                elif order_choice == "3":
                    break
                else:
                    print("Неверный выбор порядка сортировки. Пожалуйста, выберите 1, 2 или 3.")
            pass
        elif sort_choice == "3":
            while True:
                print("\nСортировка по логину:")
                print("1. По возрастанию")
                print("2. По убыванию")
                print("3. Назад в меню сортировки")
                order_choice = input("Выберите порядок сортировки: ")

                if order_choice == "1":
                    data.sort(key=lambda user: user["login"])
                    print_sorted_login(data)
                    pass
                elif order_choice == "2":
                    data.sort(key=lambda user: user["login"], reverse=True)
                    print_sorted_login(data)
                    pass
                elif order_choice == "3":
                    break
                else:
                    print("Неверный выбор порядка сортировки. Пожалуйста, выберите 1, 2 или 3.")
            pass
        elif sort_choice == "4":
            break
        else:
            print("Неверный выбор сортировки. Пожалуйста, выберите 1, 2, 3 или 4.")

def filter_active_users(data):
    print("\033[96mАктивные пользователи:\033[0m")
    active_users = [user for user in data if user['status'] == 'active']
    if not active_users:
        print("Нет активных пользователей.")
    else:
        display_users1(active_users)

def filter_inactive_users(data):
    print("\033[96mНеактивные пользователи:\033[0m")
    inactive_users = [user for user in data if user['status'] == 'inactive']
    if not inactive_users:
        print("Нет неактивных пользователей.")
    else:
        display_users1(inactive_users)


def filter_admins(data):
    print("\033[96mАдминистраторы:\033[0m")
    admins = [user for user in data if user['role'] == 1]
    if not admins:
        print("Нет администраторов.")
    else:
        display_users2(admins)

def filter_users(data):
    print("\033[96mПользователи:\033[0m")
    admins = [user for user in data if user['role'] == 0]
    if not admins:
        print("Нет пользователей.")
    else:
        display_users2(admins)

# Функция для поиска данных по введенному слову
def search_data_by_word(data, search_query):
    results = []

    for user in data:
        user_info = [str(user["id"]), user["surname"], user["name"], user["login"], user["password"],
                     "Администратор" if user["role"] == 1 else "Пользователь",
                     "Включен" if user["status"] == "active" else "Отключен"]

        for field in user_info:
            if search_query in field:
                results.append(user_info)
                break

    return results


def search_by_word(data):
    search_query = input("Введите слово для поиска: ")
    results = search_data_by_word(data, search_query)

    if results:
        table = PrettyTable()
        table.field_names = ["ID", "Фамилия", "Имя", "Логин", "Пароль", "Роль", "Статус"]

        for user_info in results:
            table.add_row(user_info)

        print("Найдены пользователи, чьи атрибуты содержат введенное слово:")
        print(table)
    else:
        print("Нет совпадений.")


def display_users1(users):
    # В этой функции можно отобразить информацию о пользователях (например, используя PrettyTable)
    # Примерно так:
    table = PrettyTable()
    table.field_names = ["ID", "Фамилия", "Имя", "Логин", "Статус"]
    for user in users:
        table.add_row([user['id'], user['surname'], user['name'], user['login'], user['status']])
    print(table)

def display_users2(users):
    # В этой функции можно отобразить информацию о пользователях (например, используя PrettyTable)
    # Примерно так:
    table = PrettyTable()
    table.field_names = ["ID", "Фамилия", "Имя", "Логин", "Роль"]
    for user in users:
        role = "Администратор" if user["role"] == 1 else "Пользователь"
        table.add_row([user['id'], user['surname'], user['name'], user['login'], role])
    print(table)


def filter_menu():
    while True:
        print("\033[96mФильтрация:\033[0m")
        print("1. Активные пользователи")  # Опция для фильтрации активных пользователей
        print("2. Отключенные пользователи")  # Опция для фильтрации отключенных пользователей
        print("3. Администраторы")  # Опция для фильтрации администраторов
        print("4. Пользователи")  # Опция для фильтрации обычных пользователей
        print("5. Назад")  # Опция для выхода из меню фильтрации
        filter_choice = input("Выберите опцию фильтрации: ")

        if filter_choice == "1":
            filter_active_users(data)
            pass
        elif filter_choice == "2":
            filter_inactive_users(data)
            pass
        elif filter_choice == "3":
            filter_admins(data)
            pass
        elif filter_choice == "4":
            filter_users(data)
            pass
        elif filter_choice == "5":
            break  # Выход из меню фильтрации
        else:
            print("Неверный выбор. Пожалуйста, выберите 1, 2, 3, 4 или 5.")


# Функция для выбора сортировка/фильтрация
def sort_or_filtr():
    while True:
        print("1. Сортировка")
        print("2. Фильтрация")
        print("3. Поиск")
        print("4. Назад")
        option_choice = input("Выберите опцию: ")

        if option_choice == "1":
            sort_menu()
            pass
        elif option_choice == "2":
            filter_menu()
            pass
        elif option_choice == "3":
            search_by_word(data)
            pass
        elif option_choice == "4":
            break
        else:
            print("Неверный выбор. Пожалуйста, выберите 1, 2, 3 или 4.")
    pass

def view_stats_login(data):
    while True:
        user_login = input("Введите логин пользователя, статистику которого вы хотите посмотреть (или нажмите Enter для отмены): ")
        if not user_login:
            return

        user = next((user for user in data if user["login"] == user_login), None)

        if user:
            login_count = user.get('login_count', 0)
            print(f"Статистика входов для пользователя с логином '{user_login}':")
            print(f"Всего входов: {login_count}")
        else:
            print(f"Пользователь с логином '{user_login}' не найден.")

def input_field(prompt):
    return input(prompt)

def get_user_input():
    delivery_num = input_field("Введите номер поставки (только цифры): ")
    while not delivery_num.isdigit:
        print("Номер поставки должен содержать только цифры. Пожалуйста, введите его заново.")
        delivery_num = input_field("Введите номер поставки (только цифры): ")

    day = input_field("Введите день договора числом: ")
    while not day.isdigit() or not (1 <= int(day) <= 31):
        print("День должен быть числом от 1 до 31. Пожалуйста, введите его заново.")
        day = input_field("Введите день договора числом: ")

    month = input_field("Введите месяц договора числом: ")
    while not month.isdigit() or not (1 <= int(month) <= 12):
        print("Месяц должен быть числом от 1 до 12. Пожалуйста, введите его заново.")
        month = input_field("Введите месяц договора числом: ")
    fields = {
        "delivery_num": delivery_num,
        "day": day,
        "month": month,
        "name_company": input_field("Введите название компании поставщика: "),
        "index": input_field("Введите индекс поставщика: "),
        "region": input_field("Введите название региона: "),
        "city": input_field("Введите название города: "),
        "street": input_field("Введите название улицы: "),
        "home_num": input_field("Введите номер дома: "),
        "ogrn_num": input_field("Введите ОГРН: "),
        "inn_num": input_field("Введите ИНН: "),
        "kp_num": input_field("Введите КПП: "),
        "rs_num": input_field("Введите расчетный счет: "),
        "bank_name": input_field("Введите название банка: "),
        "ks_num": input_field("Введите корреспондентский счет: "),
        "bik_nam": input_field("Введите БИК банка: "),
        "tel": input_field("Введите телефон: "),
        "position": input_field("Введите должность: "),
        "sokr_name": input_field("Введите инициал имени: "),
        "sokr_middlename": input_field("Введите инициал отчества: "),
        "surname": input_field("Введите фамилию: ")
    }
    return fields

def save_to_excel(data):
    df = pd.DataFrame(data, index=[0])
    try:
        existing_data = pd.read_excel("user_data.xlsx")
        df = pd.concat([existing_data, df], ignore_index=True)
    except FileNotFoundError:
        pass

    df.to_excel("user_data.xlsx", index=False)

def fill_document(doc, fields):
    for paragraph in doc.paragraphs:
        for field, value in fields.items():
            placeholder = "{" + field + "}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for field, value in fields.items():
                        placeholder = "{" + field + "}"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)

def save_document(doc, filename):
    doc.save(filename)

def create_user_folder(username):
    folder_name = os.path.join(os.getcwd(), username)  # Путь к папке с именем пользователя
    if not os.path.exists(folder_name):
        os.mkdir(folder_name)  # Создаем папку, если она не существует
    return folder_name

def save_filled_document(fields, user_login):
    doc = Document("шаблон.docx")
    fill_document(doc, fields)

    current_datetime = datetime.now()
    username = user_login
    formatted_datetime = current_datetime.strftime("%Y-%d-%m %H-%M-%S")

    user_folder = create_user_folder(username)

    docx_filename = f"{username} {formatted_datetime}.docx"
    pdf_filename = f"{username} {formatted_datetime}.pdf"

    docx_path = os.path.join(user_folder, docx_filename)
    pdf_path = os.path.join(user_folder, pdf_filename)

    doc.save(docx_path)
    convert(docx_path, pdf_path)

    return docx_path


def user_menu(data, user_login):
    while True:
        print("\033[96m1.\033[0m Сформировать документ (заполнение через диалоговое окно)")
        print("\033[96m2.\033[0m Выход")
        choice = input("Выберите действие: ")

        if choice == "1":
            fields = get_user_input()
            docx_filename = save_filled_document(fields, user_login)
            print(f"Документ успешно создан и сохранен в формате .docx с именем '{docx_filename}' в папке пользователя.")
            save_to_excel(fields)
            print("Данные успешно сохранены в Excel файл.")
        elif choice == "2":
            break
        else:
            print("Неверный выбор. Пожалуйста, выберите 1 или 2.")

def create_user_folder(username):
    folder_name = os.path.join(os.getcwd(), username)  # Путь к папке с именем пользователя
    if not os.path.exists(folder_name):
        os.mkdir(folder_name)  # Создаем папку, если она не существует
    return folder_name

def show_user_activity(username):
    width = 0.35
    data = load_data()
    users = []
    login_counts = []
    logout_counts = []
    for user in data:
        if user['login'] == username:
            users.append(user['login'])
            login_counts.append(user['login_count'])
            logout_counts.append(user.get('logout_count', 0))
    if len(users) == 0:
        print(f"No data found for user: {username}")
        return

    x = np.arange(len(users))

    plt.bar(x - width / 2, login_counts, width, label='Login count')
    plt.bar(x + width / 2, logout_counts, width, label='Logout count')

    plt.xlabel('Users')
    plt.ylabel('Count')
    plt.title(f'User Login and Logout Count for {username}')
    plt.legend()

    plt.xticks(x, users)

    plt.show()

def show_diagramMonth():
    data = load_data()

    month = int(input("Введите номер месяца (1-12): "))

    login_counts = []
    user_logins = []

    for user_data in data:
        if 'last_login' in user_data:
            login_date = datetime.datetime.strptime(user_data['last_login'], "%Y-%m-%d %H:%M:%S")
            if login_date.month == month:
                login_counts.append(user_data.get('login_count', 0))
                user_logins.append(user_data['login'])

    x = range(len(login_counts))

    plt.bar(x, login_counts, label='Входы')

    plt.xlabel('Пользователи')
    plt.ylabel('Количество входов')
    plt.title(f'График входов пользователей в выбранном месяце')
    plt.xticks(x, user_logins, rotation=90)

    plt.show()

def Graphics():
    width = 0.35
    data = load_data()
    users = []
    login_counts = []
    logout_counts = []
    for user in data:
        users.append(user['login'])
        login_counts.append(user['login_count'])
        logout_counts.append(user.get('logout_count', 0))  # Используем get для обратной совместимости
    x = np.arange(len(users))

    plt.bar(x - width / 2, login_counts, width, label='Login count')
    plt.bar(x + width / 2, logout_counts, width, label='Logout count')

    plt.xlabel('Users')
    plt.ylabel('Count')
    plt.title('User Login and Logout Count by Week')
    plt.legend()

    plt.xticks(x, users)
    plt.show()

def on_exit(login_enter):
    data = load_data()
    for user in data:
        if user["login"] == login_enter:
            user["last_exit"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            break
    save_data(data)

def show_user_activityMonth(username):
    data = load_data()

    month = int(input("Введите номер месяца (1-12): "))

    user_data = None
    for data_row in data:
        if data_row['login'] == username:
            user_data = data_row
            break

    if user_data is None:
        print(f"Пользователь с логином '{username}' не найден.")
        return

    login_counts = [0] * calendar.monthrange(datetime.datetime.now().year, month)[1]
    logout_counts = [0] * calendar.monthrange(datetime.datetime.now().year, month)[1]

    # Check if 'last_login' key exists in user_data
    if 'last_login' in user_data and user_data['last_login']:
        login_date = datetime.datetime.strptime(user_data['last_login'], "%Y-%m-%d %H:%M:%S")
        if login_date.month == month:
            login_counts[login_date.day - 1] = user_data['login_count']

    # Check if 'last_exit' key exists in user_data
    if 'last_exit' in user_data and user_data['last_exit']:
        logout_date = datetime.datetime.strptime(user_data['last_exit'], "%Y-%m-%d %H:%M:%S")
        if logout_date.month == month:
            logout_counts[logout_date.day - 1] = user_data['logout_count']

    x = list(range(1, len(login_counts) + 1))
    width = 0.4

    plt.bar(x, login_counts, width=width, label='Входы', color='blue')
    plt.bar([i + width for i in x], logout_counts, width=width, label='Выходы', color='red', alpha=0.5)

    plt.xlabel('Дни')
    plt.ylabel('Количество')
    plt.title(f'График активности пользователя {username} в выбранном месяце')
    plt.xticks([i + width / 2 for i in x], [str(day) for day in x])

    plt.legend()
    plt.show()


# Функция для админ меню
def admin_menu(data, user_login):
    while True:
        print("\033[96m1.\033[0m Посмотреть данные пользователей")
        print("\033[96m2.\033[0m Добавить нового пользователя")
        print("\033[96m3.\033[0m Удалить пользователя")
        print("\033[96m4.\033[0m Изменить данные пользователя")
        print("\033[96m5.\033[0m Изменить статус пользователя (включен/отключен)")
        print("\033[96m6.\033[0m Просмотр статистики")
        print("7. Графики")
        print("8. Графики за месяц")
        print("9. Графики за неделю конкретного пользователя")
        print("10. Графики за месяц конкретного пользователя")
        print("\033[96m11.\033[0m Выход")
        choice = input("Выберите действие: ").replace(" ", "")

        if choice == "1":
            view_users(data)
            sort_or_filter()
        elif choice == "2":
            add_user(data)
        elif choice == "3":
            delete_user(data)
        elif choice == "4":
            change_user_data(data)
        elif choice == "5":
            change_user_status(data)
        elif choice == "6":
            view_stats_menu(data)
        elif choice == "7":
            Graphics()
        elif choice == "8":
            show_diagramMonth()
        elif choice == "9":
            select = input("Выберите какую диаграмму вывести:\n1. Всех пользователей\n2. Конкретного пользователя\n")
            if select == "1":
                Graphics()
                atexit.register(on_exit, user_login)
            elif select == "2":
                username = input("Введите логин пользователя: ")
                view_users(data)
                show_user_activity(username)
                atexit.register(on_exit, user_login)
        elif choice == "10":
            select = input("Выберите какую диаграмму вывести:\n1. Всех пользователей\n2. Конкретного пользователя\n")
            if select == "1":
                show_diagramMonth()
                atexit.register(on_exit, user_login)
            elif select == "2":
                username = input("Введите логин пользователя: ")
                view_users(data)
                show_user_activityMonth(username)
                atexit.register(on_exit, user_login)
            else:
                print("Неверный выбор. Выберите 1 или 2")
        elif choice == "11":
            break
        else:
            print("Неверный выбор. Пожалуйста, выберите правильное действие.")

# Функция для вывода информации о времени продолжительности работы пользователя
def view_work_duration(user_to_view_duration):
    if user_to_view_duration:
        login_time = user_to_view_duration.get("login_time", None)
        if login_time:
            login_time = datetime.strptime(login_time, "%Y-%m-%d %H:%M:%S")
            current_time = datetime.now()
            logout_time = user_to_view_duration.get("logout_time")
            if logout_time:
                logout_time = datetime.strptime(logout_time, "%Y-%m-%d %H:%M:%S")
            else:
                logout_time = current_time
            duration = logout_time - login_time
            hours, remainder = divmod(duration.total_seconds(), 3600)
            minutes, seconds = divmod(remainder, 60)
            duration_str = f"{int(hours):02d}:{int(minutes):02d}:{int(seconds):02d}"
            print(f"Пользователь с логином '{user_to_view_duration['login']}' вошел в {login_time}, вышел в {logout_time}, и он был в системе {duration_str}.")
        else:
            print(f"Пользователь с логином '{user_to_view_duration['login']}' еще не входил в систему.")
    else:
        print("Пользователь не найден.")

# Функция для показа меню с выбором типа статистики
def view_stats_menu(data):
    while True:
        print("\033[96m1.\033[0m По входам в систему")
        print("\033[96m2.\033[0m По продолжительности работы пользователей")
        print("\033[96m3.\033[0m Выход")
        choice = input("Выберите тип статистики: ")

        if choice == "1":
            view_users(data)
            view_stats_login(data)
            pass
        elif choice == "2":
            view_users(data)
            login_to_view_duration = input("Введите логин пользователя для просмотра продолжительности работы (или нажмите Enter для отмены): ")
            user_to_view_duration = next((user for user in data if user["login"] == login_to_view_duration), None)
            view_work_duration(user_to_view_duration)
            pass
        elif choice == "3":
            break
        else:
            print("Неверный выбор. Пожалуйста, выберите 1, 2 или 3.")

# Основной цикл программы
while True:
    print("1. Вход")
    print("2. Выход")
    choice = input("Выберите действие: ")

    if choice == "1":
        data = load_data()
        authenticate(data)
    elif choice == "2":
        break
    else:
        print("Неверный выбор. Пожалуйста, выберите 1 или 2.")